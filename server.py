from http.server import HTTPServer, BaseHTTPRequestHandler
import json, requests, os
from datetime import datetime
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import tempfile
import traceback

# Verifică variabilele de mediu
KEY = os.environ.get("OPENROUTER_KEY", "").strip()
REPLICATE_KEY = os.environ.get("REPLICATE_API_TOKEN", "").strip()

print(f"=== DIAGNOSTIC ===")
print(f"OPENROUTER_KEY setat: {bool(KEY)}")
print(f"REPLICATE_API_TOKEN setat: {bool(REPLICATE_KEY)}")
if KEY:
    print(f"KEY primi 20 caractere: {KEY[:20]}...")
else:
    print("KEY: N/A")
print(f"=================")

if not KEY:
    print("⚠️ ATENȚIE: OPENROUTER_KEY nu este setată! Chatul nu va funcționa.")

# Încearcă să importe replicate
try:
    import replicate
    REPLICATE_AVAILABLE = bool(REPLICATE_KEY)
    print(f"✅ Replicate importat cu succes")
except ImportError:
    replicate = None
    REPLICATE_AVAILABLE = False
    print(f"⚠️ Replicate nu este instalat")

AGENTI = {
    "voluntari": {
        "nume": "Agent Voluntari & Outreach",
        "rol": "Esti expert in recrutare voluntari, comunicare cu comunitatea si coordonare echipe pentru ONG-uri caritabile."
    },
    "fundraising": {
        "nume": "Agent Fundraising",
        "rol": "Esti expert in strangere de fonduri pentru ONG-uri. Ajuti cu strategii fundraising, propuneri sponsorizare, campanii donatii."
    },
    "social_media": {
        "nume": "Agent Content & Social Media",
        "rol": "Esti expert in comunicare digitala pentru ONG-uri. Ajuti cu postari Facebook, Instagram, comunicate presa, newsletter."
    },
    "parteneriate": {
        "nume": "Agent Parteneriate",
        "rol": "Esti expert in parteneriate pentru ONG-uri. Ajuti cu colaborari DGASPC, primarii, firme sponsor, alte ONG-uri."
    },
    "proiecte": {
        "nume": "Agent Project & Event Manager",
        "rol": "Esti expert in managementul proiectelor si evenimentelor caritabile. Ajuti cu planificare, sarcini, termene."
    },
    "documente": {
        "nume": "Agent Rapoarte & Documente",
        "rol": "Esti expert in documentatie pentru ONG-uri. Ajuti cu rapoarte finantatori, documente ANAF, procese verbale."
    },
    "legal": {
        "nume": "Agent Legal & Conformitate",
        "rol": "Esti expert in legislatia ONG-urilor din Romania. Ajuti cu legislatie, contracte voluntari, GDPR."
    },
    "monitorizare": {
        "nume": "Agent Monitorizare & Impact",
        "rol": "Esti expert in masurarea impactului social. Ajuti cu statistici, rapoarte impact, indicatori performanta."
    },
    "comunicare": {
        "nume": "Agent Comunicare & PR",
        "rol": "Esti expert in relatii publice pentru ONG-uri. Ajuti cu comunicate presa, media, scrisori oficiale."
    }
}

INSTRUCTIUNI_ROMANA = "Raspunde EXCLUSIV in limba romana corecta, cu diacritice. Fii concis, clar si prietenos. Maximum 5 propozitii."

# Pentru Render, folosim un fișier temporar pentru istoric
FISIER_ISTORIC = "/tmp/istoric.json"

conversatie = []

def salveaza_istoric(intrebare, raspuns, agent_nume):
    try:
        if os.path.exists(FISIER_ISTORIC):
            with open(FISIER_ISTORIC, "r", encoding="utf-8") as f:
                istoric = json.load(f)
        else:
            istoric = []
        istoric.append({
            "data": datetime.now().strftime("%d.%m.%Y %H:%M"),
            "intrebare": intrebare,
            "raspuns": raspuns,
            "agent": agent_nume
        })
        if len(istoric) > 100:
            istoric = istoric[-100:]
        with open(FISIER_ISTORIC, "w", encoding="utf-8") as f:
            json.dump(istoric, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Eroare salvare istoric: {e}")

def citeste_istoric():
    try:
        if os.path.exists(FISIER_ISTORIC):
            with open(FISIER_ISTORIC, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []

def asigura_romana(text):
    if not text:
        return text
    romana_ok = any(c in text.lower() for c in ["ă", "â", "î", "ș", "ț", "și", "să", "pentru", "care", "dar"])
    if not romana_ok and len(text) > 15:
        return "⚠️ Răspunsul a fost generat tehnic. Te rog să reformulezi întrebarea pentru limba română."
    return text

def apeleaza_agent(rol_agent, mesaj, cu_istoric=False):
    if not KEY:
        return "Eroare: Cheia OPENROUTER_KEY nu este configurată. Adaug-o în Environment Variables pe Render."
    
    instructiuni = f"""{rol_agent}

REGULI OBLIGATORII:
1. Răspunde EXCLUSIV în limba română.
2. Folosește diacritice corecte (ă, â, î, ș, ț).
3. Fii concis, maximum 4-5 propoziții.
4. Nu folosi alte limbi decât româna."""
    
    mesaje = [{"role": "system", "content": instructiuni}]
    
    if cu_istoric and len(conversatie) > 0:
        for m in conversatie[-6:]:
            mesaje.append(m)
    mesaje.append({"role": "user", "content": mesaj})
    
    try:
        r = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": "Bearer " + KEY, "Content-Type": "application/json"},
            json={
                "model": "google/gemini-2.0-flash-exp:free",
                "messages": mesaje,
                "temperature": 0.5
            },
            timeout=30
        )
        data = r.json()
        if "choices" in data:
            raspuns = data["choices"][0]["message"]["content"]
            return asigura_romana(raspuns)
        else:
            print(f"Eroare API: {data}")
            return f"Eroare API: {data.get('error', {}).get('message', 'Eroare necunoscută')}"
    except Exception as e:
        print(f"Eroare request: {e}")
        return f"Eroare de conexiune: {str(e)}"

def manager_ai(tema):
    global conversatie
    conversatie.append({"role": "user", "content": tema})
    lista_agenti = ", ".join(AGENTI.keys())
    decizie = apeleaza_agent(
        "Esti Manager AI. Raspunzi DOAR cu cheia agentului potrivit. Agenti: " + lista_agenti,
        "Cerere: " + tema
    )
    agent_ales = decizie.strip().lower().replace(" ", "_")
    agent_valid = None
    for cheie in AGENTI.keys():
        if cheie in agent_ales:
            agent_valid = cheie
            break
    if not agent_valid:
        agent_valid = "voluntari"
    agent_info = AGENTI[agent_valid]
    raspuns_agent = apeleaza_agent(agent_info["rol"], tema, cu_istoric=True)
    raspuns_final = apeleaza_agent(
        "Esti Manager AI pentru o asociatie caritabila. Prezinta raspunsul agentului clar si profesional. " + INSTRUCTIUNI_ROMANA,
        "Agentul " + agent_info["nume"] + " a raspuns: " + raspuns_agent + "\nCererea: " + tema,
        cu_istoric=True
    )
    conversatie.append({"role": "assistant", "content": raspuns_final})
    if len(conversatie) > 20:
        conversatie = conversatie[-20:]
    salveaza_istoric(tema, raspuns_final, agent_info["nume"])
    return raspuns_final, agent_info["nume"]

def genereaza_imagine(prompt):
    if not REPLICATE_AVAILABLE or not REPLICATE_KEY:
        return None, "Generarea de imagini necesită token Replicate. Adaugă REPLICATE_API_TOKEN în variabilele de mediu."
    try:
        output = replicate.run(
            "black-forest-labs/flux-schnell",
            input={
                "prompt": f"{prompt}, in Romanian style, high quality",
                "num_outputs": 1,
                "width": 768,
                "height": 768
            }
        )
        if output and len(output) > 0:
            return output[0], None
        return None, "Nu am putut genera imaginea"
    except Exception as e:
        return None, f"Eroare generare: {str(e)}"

def creeaza_document_word(continut, nume_fisier="document_generat.docx"):
    try:
        doc = Document()
        doc.add_heading('Document Generat de Manager AI', 0)
        doc.add_paragraph(f"Generat la: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph()
        paragrafe = continut.split('\n')
        for p in paragrafe:
            if p.strip():
                doc.add_paragraph(p.strip())
        temp_path = os.path.join(tempfile.gettempdir(), nume_fisier)
        doc.save(temp_path)
        with open(temp_path, 'rb') as f:
            file_data = base64.b64encode(f.read()).decode('utf-8')
        os.remove(temp_path)
        return file_data, f"{nume_fisier}"
    except Exception as e:
        return None, f"Eroare creare Word: {str(e)}"

def creeaza_pdf(continut, nume_fisier="document_generat.pdf"):
    try:
        temp_path = os.path.join(tempfile.gettempdir(), nume_fisier)
        c = canvas.Canvas(temp_path, pagesize=A4)
        width, height = A4
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, "Document Generat de Manager AI")
        c.setFont("Helvetica", 10)
        c.drawString(50, height - 70, f"Generat la: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        c.setFont("Helvetica", 11)
        y = height - 100
        linii = continut.split('\n')
        for line in linii:
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 11)
            if line.strip():
                c.drawString(50, y, line.strip()[:90])
            y -= 20
        c.save()
        with open(temp_path, 'rb') as f:
            file_data = base64.b64encode(f.read()).decode('utf-8')
        os.remove(temp_path)
        return file_data, f"{nume_fisier}"
    except Exception as e:
        return None, f"Eroare creare PDF: {str(e)}"

def genereaza_document_dupa_tema(tema, tip_document="word"):
    lista_agenti = ", ".join(AGENTI.keys())
    decizie = apeleaza_agent(
        "Esti Manager AI. Raspunzi DOAR cu cheia agentului potrivit pentru a scrie un document. Agenti: " + lista_agenti,
        "Genereaza un document despre: " + tema
    )
    agent_ales = decizie.strip().lower().replace(" ", "_")
    agent_valid = None
    for cheie in AGENTI.keys():
        if cheie in agent_ales:
            agent_valid = cheie
            break
    if not agent_valid:
        agent_valid = "documente"
    agent_info = AGENTI[agent_valid]
    instructiuni_document = f"""{agent_info['rol']}

IMPORTANT: Generează un document complet, structurat, cu titlu, capitole și paragrafe. 
Documentul trebuie să fie profesional și util pentru o organizație non-profit.
Tema: {tema}

Scrie conținutul detaliat, de minimum 10-15 propoziții."""
    
    continut = apeleaza_agent(instructiuni_document, tema, cu_istoric=False)
    
    if tip_document == "pdf":
        file_data, nume = creeaza_pdf(continut, f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
    else:
        file_data, nume = creeaza_document_word(continut, f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    
    return file_data, nume, continut

# HTML (la fel ca înainte, nu modific)
HTML = """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Manager AI - Asociatie Caritabila</title>
<style>
* { margin:0; padding:0; box-sizing:border-box; }
body { font-family: Arial; background: #1a1a2e; color: white; height: 100dvh; display: flex; flex-direction: column; font-size: 16px; }
h1 { text-align:center; padding: 12px; background: #16213e; font-size: 1em; border-bottom: 1px solid #333; }
h1 span { color: #e94560; }
h1 small { display:block; font-size:0.7em; color:#aaa; margin-top:2px; }
#tabs { display:flex; background: #16213e; border-bottom: 1px solid #333; }
#tabs button { flex:1; padding: 10px; border:none; background:none; color:#aaa; cursor:pointer; font-size:0.9em; }
#tabs button.activ { color: #e94560; border-bottom: 2px solid #e94560; }
#chat-view { flex:1; overflow-y:auto; padding: 10px; display: flex; flex-direction: column; gap: 8px; }
#istoric-view { flex:1; overflow-y:auto; padding: 10px; display:none; }
.mesaj { padding: 10px 12px; border-radius: 16px; max-width: 92%; line-height: 1.5; font-size: 0.95em; word-break: break-word; }
.tu { background: #e94560; align-self: flex-end; border-bottom-right-radius: 4px; }
.manager { background: #0f3460; align-self: flex-start; border-left: 4px solid #e94560; border-bottom-left-radius: 4px; }
.agent-tag { display:inline-block; background: #e94560; color: white; font-size: 0.65em; padding: 2px 8px; border-radius: 10px; margin-bottom: 6px; }
.nume { font-weight: bold; font-size: 0.75em; margin-bottom: 4px; opacity: 0.85; }
.istoric-item { background: #0f3460; border-radius: 12px; padding: 12px; margin-bottom: 10px; border-left: 3px solid #e94560; }
.istoric-data { font-size: 0.7em; color: #aaa; margin-bottom: 4px; }
.istoric-intrebare { font-size: 0.85em; color: #f5a623; margin-bottom: 6px; font-weight: bold; }
.istoric-raspuns { font-size: 0.85em; line-height: 1.4; }
.istoric-agent { display:inline-block; background:#333; color:#aaa; font-size:0.65em; padding: 2px 8px; border-radius: 10px; margin-top: 6px; }
#input-area { display: flex; padding: 10px; gap: 8px; background: #16213e; border-top: 1px solid #333; align-items: center; flex-wrap: wrap; }
#mesaj { flex:1; padding: 12px 15px; border-radius: 25px; border: none; background: #0f3460; color: white; font-size: 1em; outline: none; min-width: 150px; }
#mesaj::placeholder { color: #aaa; }
#trimite { padding: 12px 18px; border-radius: 25px; border: none; background: #e94560; color: white; cursor: pointer; font-size: 1em; }
#microfon { padding: 12px; border-radius: 50%; border: none; background: #0f3460; color: white; cursor: pointer; font-size: 1.2em; width: 48px; height: 48px; }
#microfon.activ { background: #e94560; animation: pulse 1s infinite; }
#atasare { padding: 12px; border-radius: 50%; border: none; background: #0f3460; color: white; cursor: pointer; font-size: 1.2em; width: 48px; height: 48px; }
.generare-buttons { display: flex; gap: 8px; margin-top: 8px; flex-wrap: wrap; }
.generare-buttons button { padding: 8px 12px; border-radius: 20px; border: none; cursor: pointer; font-size: 0.8em; }
.btn-img { background: #7ed321; color: #1a1a2e; }
.btn-word { background: #f5a623; color: #1a1a2e; }
.btn-pdf { background: #e94560; color: white; }
#generare-status { font-size: 0.7em; color: #aaa; margin-top: 5px; text-align: center; }
@keyframes pulse { 0% { transform: scale(1); } 50% { transform: scale(1.1); } 100% { transform: scale(1); } }
.loading { opacity: 0.6; font-style: italic; }
.gol { text-align:center; color:#aaa; padding: 40px; }
.generated-img { max-width: 300px; border-radius: 10px; margin-top: 8px; }
</style>
</head>
<body>
<h1>🤖 <span>Manager AI</span><small>Asociatie Caritabila • 9 Agenti Specializati</small></h1>
<div id="tabs">
  <button class="activ" onclick="aratTab('chat')">💬 Chat</button>
  <button onclick="aratTab('istoric')">📋 Istoric</button>
</div>
<div id="chat-view">
  <div class="mesaj manager">
    <div class="nume">🧠 Manager AI</div>
    Salut! Sunt Manager-ul tau AI. Am 9 agenti specializati si pot genera imagini si documente! Cum te pot ajuta?
  </div>
</div>
<div id="istoric-view"></div>
<div id="input-area">
  <input type="file" id="fisier" accept="image/*,.pdf,.doc,.docx,.txt,.xls,.xlsx" style="display:none" onchange="trimiseFisier()">
  <button id="atasare" onclick="document.getElementById('fisier').click()">📎</button>
  <button id="microfon" onclick="toggleVoice()">🎤</button>
  <input id="mesaj" type="text" placeholder="Scrie sau vorbeste..." onkeypress="if(event.key==='Enter') trimite()">
  <button id="trimite" onclick="trimite()">Trimite</button>
</div>
<div class="generare-buttons">
  <button class="btn-img" onclick="genereazaImagine()">🎨 Generează Imagine</button>
  <button class="btn-word" onclick="genereazaDocument('word')">📝 Generează Word</button>
  <button class="btn-pdf" onclick="genereazaDocument('pdf')">📄 Generează PDF</button>
</div>
<div id="generare-status"></div>
<script>
let recunoastere = null;
let vorbeste = false;

function aratTab(tab) {
  document.querySelectorAll('#tabs button').forEach(b => b.classList.remove('activ'));
  event.target.classList.add('activ');
  if (tab === 'chat') {
    document.getElementById('chat-view').style.display = 'flex';
    document.getElementById('istoric-view').style.display = 'none';
  } else {
    document.getElementById('chat-view').style.display = 'none';
    document.getElementById('istoric-view').style.display = 'block';
    incarcaIstoric();
  }
}

async function incarcaIstoric() {
  const r = await fetch('/istoric');
  const data = await r.json();
  const div = document.getElementById('istoric-view');
  if (data.length === 0) {
    div.innerHTML = '<div class="gol">Nu exista conversatii salvate inca.</div>';
    return;
  }
  div.innerHTML = data.slice().reverse().map(item =>
    '<div class="istoric-item">' +
    '<div class="istoric-data">📅 ' + item.data + '</div>' +
    '<div class="istoric-intrebare">❓ ' + item.intrebare + '</div>' +
    '<div class="istoric-raspuns">' + item.raspuns + '</div>' +
    '<span class="istoric-agent">' + item.agent + '</span>' +
    '</div>'
  ).join('');
}

async function trimiseFisier() {
  const fisier = document.getElementById('fisier').files[0];
  if (!fisier) return;
  const input = document.getElementById('mesaj');
  input.value = '📎 ' + fisier.name + ' — ';
  input.focus();
  window.fisierSelectat = fisier;
}

function initVoice() {
  if (!('webkitSpeechRecognition' in window) && !('SpeechRecognition' in window)) {
    document.getElementById('microfon').style.display = 'none';
    return;
  }
  const SR = window.SpeechRecognition || window.webkitSpeechRecognition;
  recunoastere = new SR();
  recunoastere.lang = 'ro-RO';
  recunoastere.continuous = false;
  recunoastere.interimResults = false;
  recunoastere.onresult = function(e) {
    const text = e.results[0][0].transcript;
    document.getElementById('mesaj').value = text;
    toggleVoice();
    trimite();
  };
  recunoastere.onend = function() {
    vorbeste = false;
    document.getElementById('microfon').classList.remove('activ');
    document.getElementById('microfon').textContent = '🎤';
  };
}

function toggleVoice() {
  if (!recunoastere) return;
  if (vorbeste) {
    recunoastere.stop();
    vorbeste = false;
    document.getElementById('microfon').classList.remove('activ');
    document.getElementById('microfon').textContent = '🎤';
  } else {
    recunoastere.start();
    vorbeste = true;
    document.getElementById('microfon').classList.add('activ');
    document.getElementById('microfon').textContent = '🔴';
  }
}

function vorbireText(text) {
  if (!('speechSynthesis' in window)) return;
  window.speechSynthesis.cancel();
  const u = new SpeechSynthesisUtterance(text);
  u.lang = 'ro-RO';
  u.rate = 1;
  u.pitch = 1;
  window.speechSynthesis.speak(u);
}

function adauga(clasa, nume, text) {
  const chat = document.getElementById('chat-view');
  const div = document.createElement('div');
  div.className = 'mesaj ' + clasa;
  div.innerHTML = '<div class="nume">' + nume + '</div>' + text;
  chat.appendChild(div);
  chat.scrollTop = chat.scrollHeight;
  return div;
}

async function trimite() {
  const input = document.getElementById('mesaj');
  const tema = input.value.trim();
  if (!tema) return;
  input.value = '';
  adauga('tu', 'Tu', tema);
  const m = adauga('manager', '🧠 Manager AI', '<span class="loading">deleghez catre agentul potrivit...</span>');

  if (window.fisierSelectat) {
    const fisier = window.fisierSelectat;
    window.fisierSelectat = null;
    const reader = new FileReader();
    reader.onload = async function(e) {
      const base64 = e.target.result.split(',')[1];
      const tip = fisier.type;
      const r = await fetch('/chat-imagine', {
        method: 'POST',
        headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({tema: tema, base64: base64, tip: tip})
      });
      const data = await r.json();
      m.innerHTML = '<div class="nume">🧠 Manager AI</div><span class="agent-tag">' + data.agent + '</span><br>' + data.raspuns;
      vorbireText(data.raspuns);
    };
    reader.readAsDataURL(fisier);
    return;
  }

  try {
    const r = await fetch('/chat', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({tema: tema})
    });
    const data = await r.json();
    if (data.error) {
      m.innerHTML = '<div class="nume">🧠 Manager AI</div><span class="agent-tag">⚠️ Eroare</span><br>' + data.error;
    } else {
      m.innerHTML = '<div class="nume">🧠 Manager AI</div><span class="agent-tag">' + data.agent + '</span><br>' + data.raspuns;
      vorbireText(data.raspuns);
    }
  } catch(e) {
    m.innerHTML = '<div class="nume">🧠 Manager AI</div><span class="agent-tag">⚠️ Eroare</span><br>Eroare de conexiune: ' + e.message;
  }
}

async function genereazaImagine() {
  const input = document.getElementById('mesaj');
  const prompt = input.value.trim();
  if (!prompt) {
    alert('Scrie ce imagine vrei să generez!');
    return;
  }
  document.getElementById('generare-status').innerHTML = '🎨 Generez imaginea...';
  try {
    const r = await fetch('/genereaza-imagine', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({prompt: prompt})
    });
    const data = await r.json();
    if (data.success && data.url) {
      adauga('manager', '🎨 Generator AI', <img src="${data.url}" class="generated-img">);
      document.getElementById('generare-status').innerHTML = '✅ Imagine generată!';
    } else {
      document.getElementById('generare-status').innerHTML = '❌ Eroare: ' + data.error;
    }
  } catch(e) {
    document.getElementById('generare-status').innerHTML = '❌ Eroare: ' + e.message;
  }
}

async function genereazaDocument(tip) {
  const input = document.getElementById('mesaj');
  const tema = input.value.trim();
  if (!tema) {
    alert('Scrie ce document vrei să generez!');
    return;
  }
  document.getElementById('generare-status').innerHTML =📝 Generez document ${tip.toUpperCase()}...`;
  try {
    const r = await fetch('/genereaza-document', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({tema: tema, tip: tip})
    });
    const data = await r.json();
    if (data.success && data.file_data) {
      const link = document.createElement('a');
      link.href = 'data:application/octet-stream;base64,' + data.file_data;
      link.download = data.filename;
      link.click();
      
      adauga('manager', '📄 Generator Document', Am generat documentul "${data.filename}" despre: ${tema.substring(0, 100)}...);
      document.getElementById('generare-status').innerHTML = ✅ Document ${tip.toUpperCase()} generat și descărcat!;
    } else {
      document.getElementById('generare-status').innerHTML = '❌ Eroare: ' + data.error;
    }
  } catch(e) {
    document.getElementById('generare-status').innerHTML = '❌ Eroare: ' + e.message;
  }
}

initVoice();
</script>
</body>
</html>"""

class Handler(BaseHTTPRequestHandler):
    def log_message(self, format, *args): pass

    def do_GET(self):
        if self.path == '/istoric':
            istoric = citeste_istoric()
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps(istoric, ensure_ascii=False).encode())
            return
        self.send_response(200)
        self.send_header('Content-type', 'text/html; charset=utf-8')
        self.end_headers()
        self.wfile.write(HTML.encode())

    def do_POST(self):
        print(f"📨 POST request la {self.path}")
        
        if self.path == '/chat-imagine':
            try:
                data = json.loads(self.rfile.read(int(self.headers['Content-Length'])))
                tema = data['tema']
                base64_img = data['base64']
                tip = data['tip']
                r = requests.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers={"Authorization": "Bearer " + KEY, "Content-Type": "application/json"},
                    json={
                        "model": "google/gemini-2.0-flash-exp:free",
                        "messages": [
                            {"role": "user", "content": [
                                {"type": "image_url", "image_url": {"url": "data:" + tip + ";base64," + base64_img}},
                                {"type": "text", "text": "Esti Manager AI. " + tema + ". Raspunde DOAR in limba romana cu diacritice."}
                            ]}
                        ]
                    },
                    timeout=30
                )
                result = r.json()
                if "choices" in result:
                    raspuns = result["choices"][0]["message"]["content"]
                    raspuns = asigura_romana(raspuns)
                else:
                    raspuns = "Eroare la procesarea imaginii. Te rog să încerci din nou."
                salveaza_istoric(tema, raspuns, "Agent Vizual")
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"raspuns": raspuns, "agent": "Agent Vizual"}, ensure_ascii=False).encode())
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"error": str(e)}).encode())
            return

        if self.path == '/genereaza-imagine':
            try:
                data = json.loads(self.rfile.read(int(self.headers['Content-Length'])))
                prompt = data.get('prompt', '')
                url, eroare = genereaza_imagine(prompt)
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                if url:
                    self.wfile.write(json.dumps({"url": url, "success": True}).encode())
                else:
                    self.wfile.write(json.dumps({"error": eroare, "success": False}).encode())
            except Exception as e:
                self.send_response(500)
                self.end_headers()
                self.wfile.write(json.dumps({"error": str(e), "success": False}).encode())
            return

        if self.path == '/genereaza-document':
            try:
                data = json.loads(self.rfile.read(int(self.headers['Content-Length'])))
                tema = data.get('tema', '')
                tip = data.get('tip', 'word')
                file_data, nume_fisier, continut = genereaza_document_dupa_tema(tema, tip)
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                if file_data:
                    self.wfile.write(json.dumps({
                        "success": True, 
                        "file_data": file_data, 
                        "filename": nume_fisier,
                        "continut": continut
                    }).encode())
                else:
                    self.wfile.write(json.dumps({"success": False, "error": nume_fisier}).encode())
            except Exception as e:
                self.send_response(500)
                self.end_headers()
                self.wfile.write(json.dumps({"success": False, "error": str(e)}).encode())
            return

        # Endpoint-ul principal pentru chat
        if self.path == '/chat':
            try:
                data = json.loads(self.rfile.read(int(self.headers['Content-Length'])))
                tema = data['tema']
                print(f"📝 Procesez: {tema}")
                raspuns, agent_nume = manager_ai(tema)
                print(f"✅ Răspuns generat de {agent_nume}")
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"raspuns": raspuns, "agent": agent_nume}, ensure_ascii=False).encode())
            except Exception as e:
                print(f"❌ Eroare: {e}")
                traceback.print_exc()
                self.send_response(500)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({"error": str(e)}).encode())
            return

if _name_ == "_main_":
    port = int(os.environ.get("PORT", 8080))
    print(f"🚀 Aplicatia porneste pe Render...")
    print(f"🌐 Serverul rulează pe portul {port}")
    HTTPServer(('0.0.0.0', port), Handler).serve_forever()
